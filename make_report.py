from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Поля
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(1.5)

def centered(text, bold=False, size=14, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    if color:
        r.font.color.rgb = color
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

def add_figure(path, cap_text, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    if __import__('os').path.exists(path):
        run.add_picture(path, width=Cm(width_cm))
    else:
        run.text = f'[{path} — запустите блокнот]'
        run.italic = True
    caption(cap_text)

# ── ФИО ─────────────────────────────────────────────────────────────────────
centered('Притчин Андрей Андреевич', bold=True, size=16)
centered('Практическое задание №2', size=13,
         color=RGBColor(0x55, 0x55, 0x55))

doc.add_paragraph()

# ── Схема визуализации ───────────────────────────────────────────────────────
add_figure(
    'chernoff_schema.png',
    'Рис. 1. Схема визуализации лиц Чернова: '
    'черты лица при минимальных, средних и максимальных значениях параметров. '
    'Размер глаз ← REHEAT COIL Power; Наклон бровей ← VAV REHEAT Damper Position; '
    'Кривизна рта ← Thermostat Temp; Зрачки ← SUPPLY INLET Mass Flow Rate; '
    'Ширина лица ← SUPPLY INLET Temperature; Волосы ← CO₂ Concentration; '
    'Нос ← Equipment Power; Рот ← Lights Power.',
    width_cm=14
)

# ── Основная визуализация ────────────────────────────────────────────────────
add_figure(
    'chernoff_calendar.png',
    'Рис. 2. Лица Чернова для зоны F_2_Z_2 за 14 дней (31 мая – 13 июня 2016 г.). '
    'Каждая ячейка — один день. Аномальные дни: 05.06, 06.06, 08.06, 11.06, 12.06.',
    width_cm=16
)

# ── Ссылка на репозиторий ────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('Репозиторий: ')
r1.font.size = Pt(12)
r1.font.name = 'Times New Roman'
r2 = p.add_run('https://github.com/kernsedartes/visualisation')
r2.font.size = Pt(12)
r2.font.name = 'Courier New'
r2.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

doc.save('report_final.docx')
print('Готово: report_final.docx')
